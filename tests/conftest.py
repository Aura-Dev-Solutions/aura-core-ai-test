"""
Pytest configuration and fixtures for Aura Document Analyzer tests.
"""

import asyncio
import tempfile
from pathlib import Path
from typing import AsyncGenerator, Generator
from unittest.mock import AsyncMock, MagicMock

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from httpx import AsyncClient

from src.api.main import app
from src.core.config import Settings, get_settings


@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture
def test_settings() -> Settings:
    """Create test settings with temporary directories."""
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        return Settings(
            environment="testing",
            debug=True,
            database_url="sqlite+aiosqlite:///test.db",
            redis_url="redis://localhost:6379/1",
            upload_dir=temp_path / "uploads",
            vector_db_path=temp_path / "vector_db",
            log_level="DEBUG",
            secret_key="test-secret-key",
        )


@pytest.fixture
def override_settings(test_settings: Settings):
    """Override application settings for testing."""
    app.dependency_overrides[get_settings] = lambda: test_settings
    yield test_settings
    app.dependency_overrides.clear()


@pytest.fixture
def client(override_settings: Settings) -> Generator[TestClient, None, None]:
    """Create a test client for the FastAPI application."""
    with TestClient(app) as test_client:
        yield test_client


@pytest_asyncio.fixture
async def async_client(override_settings: Settings) -> AsyncGenerator[AsyncClient, None]:
    """Create an async test client for the FastAPI application."""
    async with AsyncClient(app=app, base_url="http://test") as async_test_client:
        yield async_test_client


@pytest.fixture
def sample_pdf_file() -> Generator[Path, None, None]:
    """Create a sample PDF file for testing."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
        # Create a minimal PDF content
        pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj

4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
72 720 Td
(Hello World) Tj
ET
endstream
endobj

xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000206 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
300
%%EOF"""
        temp_file.write(pdf_content)
        temp_file.flush()
        yield Path(temp_file.name)
        Path(temp_file.name).unlink(missing_ok=True)


@pytest.fixture
def sample_docx_file() -> Generator[Path, None, None]:
    """Create a sample DOCX file for testing."""
    from docx import Document
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test document for unit testing.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.save(temp_file.name)
        yield Path(temp_file.name)
        Path(temp_file.name).unlink(missing_ok=True)


@pytest.fixture
def sample_json_file() -> Generator[Path, None, None]:
    """Create a sample JSON file for testing."""
    import json
    
    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as temp_file:
        test_data = {
            "title": "Test Document",
            "content": "This is a test JSON document",
            "metadata": {
                "author": "Test Author",
                "created_date": "2024-01-01",
                "tags": ["test", "document", "json"]
            }
        }
        json.dump(test_data, temp_file, indent=2)
        temp_file.flush()
        yield Path(temp_file.name)
        Path(temp_file.name).unlink(missing_ok=True)


@pytest.fixture
def mock_embedding_model():
    """Mock embedding model for testing."""
    mock_model = MagicMock()
    mock_model.encode.return_value = [[0.1, 0.2, 0.3, 0.4] * 96]  # 384 dimensions
    return mock_model


@pytest.fixture
def mock_classification_model():
    """Mock classification model for testing."""
    mock_model = MagicMock()
    mock_model.predict.return_value = ["contract"]
    mock_model.predict_proba.return_value = [[0.1, 0.9, 0.0, 0.0]]
    return mock_model


@pytest.fixture
def mock_ner_model():
    """Mock NER model for testing."""
    mock_model = MagicMock()
    mock_doc = MagicMock()
    mock_doc.ents = [
        MagicMock(text="John Doe", label_="PERSON"),
        MagicMock(text="New York", label_="GPE"),
    ]
    mock_model.return_value = mock_doc
    return mock_model


@pytest.fixture
def mock_vector_db():
    """Mock vector database for testing."""
    mock_db = AsyncMock()
    mock_db.add_documents.return_value = None
    mock_db.search.return_value = [
        {"id": "doc1", "score": 0.95, "metadata": {"title": "Test Doc 1"}},
        {"id": "doc2", "score": 0.85, "metadata": {"title": "Test Doc 2"}},
    ]
    return mock_db


@pytest.fixture
def mock_database():
    """Mock database for testing."""
    mock_db = AsyncMock()
    return mock_db


# Markers for different test types
pytest.mark.unit = pytest.mark.unit
pytest.mark.integration = pytest.mark.integration
pytest.mark.e2e = pytest.mark.e2e
pytest.mark.slow = pytest.mark.slow
