"""
DOCX document extractor for Microsoft Word documents.
"""

import asyncio
from pathlib import Path
from docx import Document as DocxDocument
from src.core.exceptions import DocumentProcessingError
from src.core.models import DocumentMetadata, DocumentType
from .base import BaseDocumentExtractor


class DOCXExtractor(BaseDocumentExtractor):
    """DOCX document extractor."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size, chunk_overlap)
        self.supported_types = [DocumentType.DOCX]
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        def _extract():
            try:
                doc = DocxDocument(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                return "\n\n".join(text_parts)
                
            except Exception as e:
                raise DocumentProcessingError(f"Failed to extract DOCX text: {str(e)}") from e
        
        text = await asyncio.get_event_loop().run_in_executor(None, _extract)
        return self._clean_text(text)
    
    async def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from DOCX file."""
        def _extract_metadata():
            try:
                doc = DocxDocument(file_path)
                props = doc.core_properties
                
                # Count words
                word_count = 0
                for paragraph in doc.paragraphs:
                    word_count += len(paragraph.text.split())
                
                return DocumentMetadata(
                    title=props.title,
                    author=props.author,
                    created_date=props.created,
                    modified_date=props.modified,
                    word_count=word_count,
                    file_size=file_path.stat().st_size,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    custom_fields={
                        'paragraph_count': len(doc.paragraphs),
                        'table_count': len(doc.tables),
                    }
                )
                
            except Exception as e:
                return DocumentMetadata(
                    file_size=file_path.stat().st_size,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract_metadata)
